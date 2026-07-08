#!/usr/bin/env python3
"""產生 001-借款契約書 PDF — 最終版"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def FN(r,s=12,b=False):
    r.font.name='cwTeXKai'; r.font.size=Pt(s); r.bold=b
    r.element.rPr.rFonts.set(qn('w:eastAsia'),'cwTeXKai')

doc=Document()
s=doc.sections[0]
s.page_width=Cm(21); s.page_height=Cm(29.7)
s.top_margin=Cm(2.5); s.bottom_margin=Cm(2.5)
s.left_margin=Cm(2.5); s.right_margin=Cm(2.5)
sty=doc.styles['Normal']
sty.font.name='cwTeXKai'; sty.font.size=Pt(12)
sty.paragraph_format.line_spacing=1.5
sty.element.rPr.rFonts.set(qn('w:eastAsia'),'cwTeXKai')

def C(t,s=18,b=True):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(t); FN(r,s,b); p.paragraph_format.space_after=Pt(4)
def H(t):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(10)
    r=p.add_run(t); FN(r,b=True)
def P(t,i=1.0):
    p=doc.add_paragraph(); p.paragraph_format.first_line_indent=Cm(i)
    p.paragraph_format.space_after=Pt(3); r=p.add_run(t); FN(r)
def B(l,w=40):
    p=doc.add_paragraph(); p.paragraph_format.first_line_indent=Cm(1)
    p.add_run(l+'：'); p.add_run('_'*w)

C('借款契約書',20)
P('義務人即債務人（以下簡稱甲方）因資金週轉需要，提供下列標的物辦理抵押權登記予債權人（以下簡稱乙方）作為借款擔保物，標的明細如下：')
H('不動產標示')
t=doc.add_table(rows=3,cols=5); t.style='Table Grid'
for j,h in enumerate(['標的物類型','鄉鎮市區','段/小段','地號／建號','權利範圍']):
    t.cell(0,j).text=h
    for p in t.cell(0,j).paragraphs:
        for r in p.runs: FN(r,11,b=True)
        p.alignment=WD_ALIGN_PARAGRAPH.CENTER
for j,v in enumerate(['土地','桃園市桃園區','大樹林段','19444','27/10000']):
    t.cell(1,j).text=v
for p in t.cell(1,j).paragraphs:
    for r in p.runs: FN(r,11)
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
for j,v in enumerate(['建物','桃園市桃園區','成功路二段1號8樓之5','6028','1/1']):
    t.cell(2,j).text=v
for p in t.cell(2,j).paragraphs:
    for r in p.runs: FN(r,11)
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
H('一、借款金額')
P('甲方向乙方借貸金額為新臺幣壹佰萬元整，並同意以借款金額1.5倍為最高限額抵押權設定登記予乙方以供擔保債權之清償。甲方另開立未填寫到期日之本票予乙方作為擔保債務履行之憑據，並授權乙方視實際需要自行填寫到期日及法定利率以行使票據上權利。')
H('二、借款期限')
P('借款期限自中華民國115年7月8日起至118年7月7日止（共計三年）。未有另行約定者，甲方應於屆期日全數清償，逾期同意以違約論處，加計每萬元每日違約金。甲方同意於債務清償時將本金連同違約金一次性歸還乙方。')
H('三、利息計算與清償')
P('雙方約定以每一個月為一期，利息以月息2%計算，利息起算日為每月8日。甲方應於每月8日繳納利息，逾期視為違約。若甲方於借款期間六個月後清償全部或部分借款者，不在此限。')
H('四、債務未清償之後果')
P('若債務屆期未受清償，乙方可向法院申請強制執行本擔保抵押物拍賣，拍賣價金除支付借款金額外，須另行支付執行費用、訴訟費用及律師費用等。')
H('五、抵押物之合法性')
P('甲方保證提供設定抵押之不動產完全為其合法所有，且無設定抵押權、典權、租賃權或任何足以妨礙抵押權行使之情事。如有不實，願負全部法律責任。')
H('六、費用負擔')
P('一、設定登記費：按擔保債權金額千分之一計算。^[土地法§76]')
P('二、他項證明書費：每張80元。^[地政規費標準]')
P('三、地政規費（謄本等）：依地政事務所實際收費。')
P('四、代書費：由雙方與地政士另行議定。')
P('五、印花稅：按擔保債權金額千分之一。^[印花稅法§7④]')
P('六、以上費用由______方負擔。')
H('七、管轄法院及送達')
P('因本契約所生爭議，雙方合意以臺灣桃園地方法院為第一審管轄法院。通知或催告得以書面、簡訊或LINE為之，發出即視為送達。本契約壹式參份，雙方各執壹份，地政機關或代書存查壹份。')
p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(20); C('簽署頁',14)
H('甲方（債務人）'); B('姓　　名'); B('身分證字號'); B('電　　話')
H('乙方（債權人）'); B('姓　　名'); B('身分證字號'); B('電　　話')
C('\n中華民國______年______月______日',12,False)
doc.save('/tmp/001-借款契約書.docx')
print('✅ 001 complete')
