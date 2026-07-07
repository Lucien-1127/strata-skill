#!/usr/bin/env python3
"""
generate_doc.py — 台灣法律文件一鍵生成腳本

用法:
    python3 generate_doc.py --track court --title "民事起訴狀" --content content.txt
    python3 generate_doc.py --track contract --title "借款契約" --output /tmp/contract.docx
    python3 generate_doc.py --track gov --title "開會通知單" --pdf   # docx + pdf 一鍵完成

支援三軌道：
    court   法院書狀（cwTeXKai 14pt，邊距2.5cm）
    gov     政府公文（AR PL UKai TW 12pt，左側裝訂線）
    contract 商業合約（Noto Sans CJK TC 12pt，縮排層級）
"""

import argparse, sys, os, subprocess
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


# ── 三軌道參數 ───────────────────────────────────────
TRACKS = {
    'court': {
        'font': 'cwTeXKai',
        'title_size': Pt(20),
        'body_size': Pt(14),
        'line_spacing': Pt(28),
        'margin': 2.5,
        'tab_stops': [1.5, 3.0],  # 縮排 2 字、4 字
        'style': '法院書狀',
    },
    'gov': {
        'font': 'AR PL UKai TW',
        'title_size': Pt(16),
        'body_size': Pt(12),
        'line_spacing': Pt(20),
        'margin': 2.5,
        'left_gutter': 1.5,
        'style': '政府公文',
    },
    'contract': {
        'font': 'Noto Sans CJK TC',
        'title_size': Pt(18),
        'body_size': Pt(12),
        'line_spacing': Pt(22),
        'margin': 2.0,
        'indent_levels': [2, 4, 6],  # 縮排字數：一、(一)、1.
        'style': '商業合約',
    },
}


def _apply_paragraph_protection(p, keep_next=False):
    """Apply OxmlElement keepLines + optional keepNext for orphan prevention.
    
    This mirrors the SKILL.md SOP Step 2 requirements:
    - keepLines: prevents page break inside a paragraph
    - keepNext: keeps this paragraph with the next one (for headings)
    """
    from docx.oxml import OxmlElement
    pPr = p._element.get_or_add_pPr()
    for tag in ['w:keepLines']:
        el = OxmlElement(tag)
        pPr.append(el)
    if keep_next:
        el = OxmlElement('w:keepNext')
        pPr.append(el)


def set_font(run, font_name, size=None, bold=False):
    """設定 run 的東亞字型"""
    run.font.name = font_name
    if size:
        run.font.size = size
    if bold:
        run.bold = True
    run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def make_doc(track, title, content_lines, output):
    """生成 .docx"""
    cfg = TRACKS[track]
    doc = Document()

    # 頁面設定
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    m = cfg['margin']
    sec.top_margin = Cm(m)
    sec.bottom_margin = Cm(m)
    if track == 'gov':
        sec.left_margin = Cm(m + 1.5)  # 左側裝訂線
    else:
        sec.left_margin = Cm(m)
    sec.right_margin = Cm(m)

    # 預設樣式
    style = doc.styles['Normal']
    style.font.name = cfg['font']
    style.font.size = cfg['body_size']
    style.element.rPr.rFonts.set(qn('w:eastAsia'), cfg['font'])
    style.paragraph_format.line_spacing = cfg['line_spacing']

    # 標題
    p = doc.add_paragraph(title)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.runs[0], cfg['font'], cfg['title_size'], bold=True)
    _apply_paragraph_protection(p, keep_next=True)
    p.paragraph_format.keep_with_next = True

    doc.add_paragraph('')  # 空行

    # 內文
    for line in content_lines:
        line = line.strip()
        if not line:
            doc.add_paragraph('')
            continue

        indent = 0
        text = line
        # 自動判斷縮排層級
        if line.startswith('（'):
            indent = 2  # (一)(二) 縮排4字
        elif line.startswith(('一、', '二、', '三、')):
            indent = 1  # 一、二、三 縮排2字
        elif line.startswith(('1.', '2.', '3.')):
            indent = 3  # 1. 2. 3. 縮排6字

        p = doc.add_paragraph(text)
        p.paragraph_format.widow_control = True
        _apply_paragraph_protection(p)
        if indent > 0:
            p.paragraph_format.left_indent = Cm(indent * 0.7)  # 每級約0.7cm

    doc.save(output)
    return output


def to_pdf(docx_path):
    """LibreOffice 轉 PDF"""
    pdf_path = docx_path.replace('.docx', '.pdf')
    result = subprocess.run(
        ['libreoffice', '--headless', '--convert-to', 'pdf', docx_path,
         '--outdir', os.path.dirname(docx_path)],
        capture_output=True, text=True, timeout=120)
    if os.path.exists(pdf_path):
        return pdf_path
    print(f'PDF 轉換失敗: {result.stderr}', file=sys.stderr)
    return None


def main():
    parser = argparse.ArgumentParser(description='台灣法律文件一鍵生成')
    parser.add_argument('--track', choices=['court', 'gov', 'contract'], required=True)
    parser.add_argument('--title', default='文件標題')
    parser.add_argument('--content', help='內文檔案路徑（一行一段）')
    parser.add_argument('--output', default=None, help='輸出 .docx 路徑')
    parser.add_argument('--pdf', action='store_true', help='同時轉 PDF')
    args = parser.parse_args()

    # 讀取內文
    if args.content and os.path.exists(args.content):
        with open(args.content, encoding='utf-8') as f:
            lines = f.readlines()
    else:
        lines = [sys.stdin.read()] if not sys.stdin.isatty() else ['（請輸入內文）']

    # 決定輸出路徑
    if not args.output:
        ts = os.path.splitext(os.path.basename(args.content or 'output'))[0]
        args.output = f'/tmp/{ts}_{args.track}.docx'

    # 生成
    make_doc(args.track, args.title, lines, args.output)
    print(f'✅ docx: {args.output}')

    if args.pdf:
        pdf = to_pdf(args.output)
        if pdf:
            print(f'✅ pdf:  {pdf}')

    cfg = TRACKS[args.track]
    print(f'   軌道: {cfg["style"]}')
    print(f'   字型: {cfg["font"]} {cfg["body_size"]}')


if __name__ == '__main__':
    main()
